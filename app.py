import streamlit as st
import PyPDF2
import docx
import os
import json
import time
from io import BytesIO
from PIL import Image
from google import genai
from google.genai import types
from google.genai.errors import APIError

# Para geração de PDF na Camada 6
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

st.set_page_config(page_title="Gestor Profissional de Carreira & Currículos", page_icon="💼", layout="wide")

st.markdown(
    """
    <style>
        html, body, [class*="css"] {
            translate: no;
        }
    </style>
    <meta name="google" content="notranslate">
    """,
    unsafe_allow_html=True
)

# --- DIRETÓRIOS E BANCO DE DADOS LOCAL ---
DATA_DIR = "user_data"
os.makedirs(DATA_DIR, exist_ok=True)

USERS_DB_FILE = os.path.join(DATA_DIR, "users_auth.json")

def load_users_db():
    if os.path.exists(USERS_DB_FILE):
        with open(USERS_DB_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

def save_users_db(users_db):
    with open(USERS_DB_FILE, "w", encoding="utf-8") as f:
        json.dump(users_db, f, ensure_ascii=False, indent=4)

def get_user_folder(username):
    user_dir = os.path.join(DATA_DIR, username)
    os.makedirs(user_dir, exist_ok=True)
    return user_dir

def get_profile_file(username):
    return os.path.join(get_user_folder(username), "profile_info.json")

def get_knowledge_file(username):
    return os.path.join(get_user_folder(username), "knowledge_base.txt")

def load_profile(username):
    p_file = get_profile_file(username)
    if os.path.exists(p_file):
        with open(p_file, "r", encoding="utf-8") as f:
            return json.load(f)
    return {
        "nome_completo": "",
        "email": "",
        "telefone": "",
        "endereco": "",
        "linkedin": "",
        "github": "",
        "api_key": ""
    }

def save_profile(username, profile_data):
    p_file = get_profile_file(username)
    with open(p_file, "w", encoding="utf-8") as f:
        json.dump(profile_data, f, ensure_ascii=False, indent=4)

def carregar_base_conhecimento(username):
    k_file = get_knowledge_file(username)
    if os.path.exists(k_file):
        with open(k_file, "r", encoding="utf-8") as f:
            return f.read()
    return ""

def salvar_na_base_conhecimento(username, novo_texto):
    k_file = get_knowledge_file(username)
    texto_atual = carregar_base_conhecimento(username)
    texto_atualizado = (texto_atual + "\n\n" + novo_texto).strip()
    with open(k_file, "w", encoding="utf-8") as f:
        f.write(texto_atualizado)

def limpar_base_conhecimento(username):
    k_file = get_knowledge_file(username)
    if os.path.exists(k_file):
        os.remove(k_file)

# --- FUNÇÃO RESILIENTE COM FALLBACK DE MODELOS GRATUITOS ---
def chamar_gemini_com_fallback(client, contents, max_retries=2):
    """Tenta executar nos modelos gratuitos da Google alternando se estourar o limite 429."""
    modelos_gratuitos = [
        'gemini-2.0-flash-lite',
        'gemini-2.0-flash',
        'gemini-1.5-flash'
    ]
    
    for modelo in modelos_gratuitos:
        for tentativa in range(max_retries):
            try:
                res = client.models.generate_content(
                    model=modelo,
                    contents=contents
                )
                return res.text
            except Exception as e:
                err_str = str(e)
                if "429" in err_str or "RESOURCE_EXHAUSTED" in err_str:
                    tempo_espera = 4 * (tentativa + 1)
                    st.warning(f"Cota do modelo '{modelo}' atingida. Aguardando {tempo_espera}s... (Alternando modelo automaticamente)")
                    time.sleep(tempo_espera)
                else:
                    break
                    
    raise Exception("Todos os modelos gratuitos atingiram o limite temporário. Aguarde 1 minuto e tente novamente.")

# --- EXTRAÇÃO DE ARQUIVOS (CAMADA 3) ---
def extrair_texto_arquivo(file, api_key=None):
    nome = file.name
    if nome.endswith(".docx"):
        doc = docx.Document(file)
        return f"\n--- Conteúdo do arquivo: {nome} ---\n" + "\n".join([p.text for p in doc.paragraphs])
    elif nome.endswith(".txt"):
        return f"\n--- Conteúdo do arquivo: {nome} ---\n" + file.read().decode("utf-8")
    elif nome.endswith(".pdf"):
        reader = PyPDF2.PdfReader(file)
        texto = ""
        for page in reader.pages:
            texto += page.extract_text() or ""
        return f"\n--- Conteúdo do arquivo: {nome} ---\n" + texto
    elif file.type.startswith("image/"):
        if not api_key:
            st.error(f"Insira sua API Key para processar imagens ({nome}).")
            return ""
        try:
            client = genai.Client(api_key=api_key)
            img = Image.open(file)
            img.thumbnail((800, 800))  # Reduz tamanho para economizar tokens
            buffer = BytesIO()
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            img.save(buffer, format="JPEG", quality=80)
            img_bytes = buffer.getvalue()
            
            prompt_ocr = "Extraia todo o texto, datas, títulos e informações relevantes deste documento/certificado."
            part_img = types.Part.from_bytes(data=img_bytes, mime_type="image/jpeg")
            
            texto_extraido = chamar_gemini_com_fallback(
                client=client,
                contents=[part_img, prompt_ocr]
            )
            return f"\n--- Transcrição da imagem: {nome} ---\n" + texto_extraido
        except Exception as e:
            st.error(f"Erro ao ler a imagem {nome}: {e}")
            return ""
    return ""

# --- CONVERSORES PARA EXPORTAÇÃO (CAMADA 6) ---
def gerar_docx(texto_markdown):
    doc = docx.Document()
    for linha in texto_markdown.split("\n"):
        if linha.startswith("# "):
            doc.add_heading(linha.replace("# ", ""), level=1)
        elif linha.startswith("## "):
            doc.add_heading(linha.replace("## ", ""), level=2)
        elif linha.startswith("### "):
            doc.add_heading(linha.replace("### ", ""), level=3)
        elif linha.startswith("- ") or linha.startswith("* "):
            doc.add_paragraph(linha[2:], style='List Bullet')
        else:
            doc.add_paragraph(linha)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def gerar_pdf(texto_markdown):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    story = []

    style_normal = styles['Normal']
    style_normal.fontSize = 10
    style_normal.leading = 14

    style_heading = styles['Heading1']
    style_heading.fontSize = 14
    style_heading.leading = 18

    for linha in texto_markdown.split("\n"):
        clean_line = linha.replace("**", "").replace("*", "").strip()
        if not clean_line:
            story.append(Spacer(1, 6))
            continue
        
        if linha.startswith("# ") or linha.startswith("## ") or linha.startswith("### "):
            story.append(Paragraph(clean_line.replace("#", "").strip(), style_heading))
            story.append(Spacer(1, 4))
        else:
            story.append(Paragraph(clean_line, style_normal))
            story.append(Spacer(1, 3))
            
    doc.build(story)
    buffer.seek(0)
    return buffer

# --- ESTADO DE SESSÃO DO USUÁRIO ---
if "usuario_logado" not in st.session_state:
    st.session_state["usuario_logado"] = None
if "historico_chat" not in st.session_state:
    st.session_state["historico_chat"] = []

# Chave para reset do campo de upload de arquivos e texto
if "uploader_key" not in st.session_state:
    st.session_state["uploader_key"] = 0

# --- CAMADA 2: PÁGINA DE LOGIN E CADASTRO ---
if not st.session_state["usuario_logado"]:
    st.title("💼 Gestor Profissional de Carreira & Currículos")
    
    tab_login, tab_cadastro = st.tabs(["🔒 Entrar no Sistema", "📝 Criar Nova Conta"])
    users_db = load_users_db()

    with tab_login:
        st.subheader("Acesse seu Perfil Profissional")
        user_input = st.text_input("Usuário ou Email:", key="login_user")
        pass_input = st.text_input("Senha:", type="password", key="login_pass")
        
        if st.button("Entrar", type="primary"):
            if user_input in users_db and users_db[user_input]["password"] == pass_input:
                st.session_state["usuario_logado"] = user_input
                st.success("Login efetuado com sucesso!")
                st.rerun()
            else:
                st.error("Usuário/Email ou senha incorretos.")

    with tab_cadastro:
        st.subheader("Cadastre-se para criar seus currículos")
        new_user = st.text_input("Escolha um Nome de Usuário (sem espaços):", key="cad_user")
        new_email = st.text_input("Seu Email:", key="cad_email")
        new_pass = st.text_input("Escolha uma Senha:", type="password", key="cad_pass")
        
        if st.button("Cadastrar e Acessar"):
            if not new_user or not new_pass or not new_email:
                st.warning("Preencha todos os campos obrigatórios.")
            elif new_user in users_db:
                st.error("Este nome de usuário já está em uso.")
            else:
                users_db[new_user] = {"email": new_email, "password": new_pass}
                save_users_db(users_db)
                
                p_data = load_profile(new_user)
                p_data["email"] = new_email
                save_profile(new_user, p_data)

                st.session_state["usuario_logado"] = new_user
                st.success("Conta criada com sucesso!")
                st.rerun()

    st.stop()

# --- APLICAÇÃO PRINCIPAL (USUÁRIO AUTENTICADO) ---
usuario = st.session_state["usuario_logado"]
perfil = load_profile(usuario)

# SIDEBAR: CONTROLE E LOGOUT
with st.sidebar:
    st.title(f"👤 Olá, {usuario}!")
    if st.button("🚪 Sair / Logout"):
        st.session_state["usuario_logado"] = None
        st.session_state["historico_chat"] = []
        st.rerun()
    st.markdown("---")

# TABS PRINCIPAIS
tab1, tab2, tab3, tab4 = st.tabs([
    "1️⃣ Perfil & Dados Pessoais",
    "2️⃣ Banco de Dados de Documentos",
    "3️⃣ Gerar Currículo pela Vaga",
    "4️⃣ Refinamento, Portfólio & Exportação"
])

# --- CAMADA 1: CRIAÇÃO / EDIÇÃO DE PERFIL ---
with tab1:
    st.header("1. Informações Pessoais & Trabalhistas")
    st.caption("Preencha seus dados de contato e links profissionais para compor o cabeçalho automático dos seus currículos.")

    with st.form("form_perfil"):
        col_p1, col_p2 = st.columns(2)
        with col_p1:
            nome = st.text_input("Nome Completo:", value=perfil.get("nome_completo", ""))
            email = st.text_input("Email de Contato:", value=perfil.get("email", ""))
            telefone = st.text_input("Telefone / WhatsApp:", value=perfil.get("telefone", ""))
            endereco = st.text_input("Endereço / Cidade e Estado:", value=perfil.get("endereco", ""))
        
        with col_p2:
            linkedin = st.text_input("Link do LinkedIn:", value=perfil.get("linkedin", ""))
            github = st.text_input("Link do GitHub / Portfólio:", value=perfil.get("github", ""))
            api_key = st.text_input("Sua API Key do Gemini (Armazenada com Segurança):", value=perfil.get("api_key", ""), type="password")
        
        btn_salvar_perfil = st.form_submit_button("💾 Salvar Dados do Perfil")

        if btn_salvar_perfil:
            novo_perfil = {
                "nome_completo": nome,
                "email": email,
                "telefone": telefone,
                "endereco": endereco,
                "linkedin": linkedin,
                "github": github,
                "api_key": api_key
            }
            save_profile(usuario, novo_perfil)
            st.success("Perfil atualizado com sucesso!")
            st.rerun()

# --- CAMADA 3: BANCO DE DADOS DE DOCUMENTOS ---
with tab2:
    st.header("2. Banco de Dados de Documentos & Conhecimento")
    st.caption("Suba suas experiências: Carteira de trabalho digital, declarações de trabalho informal, cursos, eventos, certificados e graduações.")

    st.subheader("➕ Adicionar Novos Documentos ao Banco")
    
    # Utiliza chave dinâmica baseada em uploader_key para permitir reset dos inputs
    texto_livre = st.text_area(
        "Adicionar texto/resumo manual de conquistas:", 
        placeholder="Ex: Trabalhei como freelancer desenvolvendo sistemas em Python entre 2022 e 2024...",
        key=f"texto_livre_{st.session_state['uploader_key']}"
    )
    
    uploaded_files = st.file_uploader(
        "Selecione arquivos (PDF, DOCX, TXT, Imagens de Certificados):", 
        type=["pdf", "txt", "docx", "png", "jpg", "jpeg", "webp"], 
        accept_multiple_files=True,
        key=f"file_uploader_{st.session_state['uploader_key']}"
    )

    if st.button("📥 Processar e Salvar no Banco de Dados"):
        if not texto_livre and not uploaded_files:
            st.warning("Insira algum texto ou selecione arquivos antes de salvar.")
        else:
            with st.spinner("Extraindo e integrando dados ao seu histórico..."):
                conteudo_adicional = ""
                if texto_livre:
                    conteudo_adicional += f"\n--- Declaração / Registro Manual ---\n{texto_livre}"
                
                if uploaded_files:
                    for f in uploaded_files:
                        conteudo_adicional += extrair_texto_arquivo(f, perfil.get("api_key"))
                        time.sleep(2)  # Pausa de 2 segundos para evitar limite de requisições 429
                
                if conteudo_adicional.strip():
                    salvar_na_base_conhecimento(usuario, conteudo_adicional)
                    st.success("Documentos gravados com sucesso no banco de dados!")
                    
                    # Incrementa a chave para limpar a caixa de upload e o texto livre da tela
                    st.session_state["uploader_key"] += 1
                    time.sleep(1)
                    st.rerun()

    st.markdown("---")
    st.subheader("📚 Consultar / Excluir Banco de Dados Salvo")
    base_salva = carregar_base_conhecimento(usuario)
    
    if base_salva:
        st.info(f"O seu banco de dados atual possui ~{len(base_salva)} caracteres de histórico compilado.")
        with st.expander("Clique para visualizar todo o conteúdo do seu banco de dados extraído"):
            st.text_area("Histórico Completo Extraído:", value=base_salva, height=250, disabled=True)
        
        if st.button("🗑️ Excluir / Resetar Todo o Banco de Dados", type="secondary"):
            limpar_base_conhecimento(usuario)
            st.success("Banco de dados resetado com sucesso.")
            st.rerun()
    else:
        st.warning("Nenhum documento salvo no seu banco de dados ainda.")

# --- CAMADA 4: INCLUSÃO DA VAGA DESEJADA ---
with tab3:
    st.header("3. Inclusão da Vaga & Escolha do Modelo")
    
    vaga_input = st.text_area("Cole a descrição completa ou os requisitos da Vaga de Emprego desejada:", height=200, placeholder="Ex: Procuramos Desenvolvedor Python com experiência em Streamlit, APIs REST...")
    
    modelo_estilo = st.selectbox(
        "Escolha o Modelo do Currículo:",
        [
            "Tech / Moderno (Focado em hard e soft skills, projetos e ferramentas)",
            "Executivo / Tradicional (Focado em cargos, metas atingidas e liderança)",
            "Clean / Minimalista (Direto e conciso para rápida leitura)",
            "Transição de Carreira (Evidencia competências transferíveis e projetos)"
        ]
    )

    if st.button("🚀 Criar Currículo Personalizado + Recomendações de Portfólio", type="primary"):
        api_key_curr = perfil.get("api_key")
        base_dados_curr = carregar_base_conhecimento(usuario)

        if not api_key_curr:
            st.error("Insira sua API Key do Gemini na aba 'Perfil & Dados Pessoais' antes de continuar.")
        elif not base_dados_curr:
            st.error("Seu Banco de Dados de Documentos está vazio. Suba certificados ou histórico na aba 2.")
        elif not vaga_input:
            st.warning("Cole a descrição da vaga para continuar.")
        else:
            try:
                with st.spinner("Consultando seu banco de dados, alinhando com a vaga e analisando portfólio..."):
                    client = genai.Client(api_key=api_key_curr)

                    info_pessoais = f"""
                    NOME COMPLETO: {perfil.get('nome_completo')}
                    EMAIL: {perfil.get('email')}
                    TELEFONE: {perfil.get('telefone')}
                    ENDEREÇO: {perfil.get('endereco')}
                    LINKEDIN: {perfil.get('linkedin')}
                    GITHUB/PORTFÓLIO: {perfil.get('github')}
                    """

                    prompt_geracao = f"""
                    Você é um Consultor Profissional de Carreira, especialista em RH e mentor técnico de Portfólio.

                    REGRAS ABSOLUTAS DO CURRÍCULO:
                    1. Zero Inventividade na Seção de Currículo: Utilize APENAS as informações reais presentes no Banco de Dados do Usuário.
                    2. Estilo Selecionado: [{modelo_estilo}].
                    3. Formate a saída em Markdown limpo e estruturado.

                    --- DADOS DE CABEÇALHO DO USUÁRIO ---
                    {info_pessoais}

                    --- BANCO DE DADOS DE DOCUMENTOS E EXPERIÊNCIAS ---
                    {base_dados_curr}

                    --- VAGA ALVO ---
                    {vaga_input}

                    --- ESTRUTURA REQUERIDA ---
                    (PARTE 1: CURRÍCULO)
                    1. Cabeçalho Completo
                    2. Resumo Profissional (Direcionado para a vaga alvo)
                    3. Principais Competências & Habilidades Táticas
                    4. Experiências Profissionais / Histórico de Trabalho
                    5. Formação Acadêmica, Cursos e Certificados Extraídos

                    ---
                    (PARTE 2: ACONSELHAMENTO DE PORTFÓLIO & LACUNAS TÉCNICAS)
                    Adicione uma seção final destacada chamada:
                    "🛠️ PLANO DE ADEQUAÇÃO DE PORTFÓLIO E RECOMENDAÇÕES DE PROJETOS"
                    
                    Analise os requisitos da VAGA em relação ao BANCO DE DADOS do usuário e forneça:
                    - **Análise de Aderência (% estimada de compatibilidade)**.
                    - **Lacunas Identificadas**: Quais ferramentas, métodos ou conhecimentos solicitados pela vaga o candidato ainda não possui no histórico?
                    - **Projetos Recomendados para Construir**: Sugira de 1 a 3 projetos práticos reais (passo a passo resumido, tecnologias sugeridas) que o usuário pode desenvolver no GitHub/Portfólio para suprir exatamente essas lacunas e se destacar no processo seletivo.
                    """

                    texto_gerado = chamar_gemini_com_fallback(
                        client=client,
                        contents=prompt_geracao
                    )

                    st.session_state["historico_chat"] = [
                        {"role": "assistant", "content": texto_gerado}
                    ]
                    st.success("Currículo e Análise de Portfólio gerados! Vá para a aba 'Refinamento, Portfólio & Exportação'.")

            except Exception as e:
                st.error(f"Erro ao processar a requisição: {e}")

# --- CAMADA 5 E 6: REFINAMENTO & EXPORTAÇÃO ---
with tab4:
    st.header("4. Refinamento, Portfólio & Exportação")
    
    if not st.session_state["historico_chat"]:
        st.info("Nenhum currículo foi gerado ainda. Preencha os requisitos na aba 'Gerar Currículo pela Vaga'.")
    else:
        curriculo_atual = st.session_state["historico_chat"][-1]["content"]

        # --- CAMADA 6: EXPORTAÇÃO ---
        st.subheader("📥 Baixar Currículo no Formato Desejado")
        col_down1, col_down2, col_down3 = st.columns(3)

        with col_down1:
            st.download_button(
                label="📄 Baixar em Markdown / TXT",
                data=curriculo_atual,
                file_name=f"curriculo_{usuario}.md",
                mime="text/markdown"
            )

        with col_down2:
            try:
                docx_buffer = gerar_docx(curriculo_atual)
                st.download_button(
                    label="📝 Baixar em DOCX (Word)",
                    data=docx_buffer,
                    file_name=f"curriculo_{usuario}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"Erro ao gerar DOCX: {e}")

        with col_down3:
            try:
                pdf_buffer = gerar_pdf(curriculo_atual)
                st.download_button(
                    label="🔴 Baixar em PDF",
                    data=pdf_buffer,
                    file_name=f"curriculo_{usuario}.pdf",
                    mime="application/pdf"
                )
            except Exception as e:
                st.error(f"Erro ao gerar PDF: {e}")

        st.markdown("---")
        
        # --- CAMADA 5: REFINAMENTO VIA CHAT ---
        st.subheader("💬 Solicitar Alterações, Dúvidas de Projetos ou Ajustes Finais")
        st.caption("Você pode pedir alterações no currículo ou tirar dúvidas sobre como desenvolver os projetos sugeridos para o portfólio.")

        for msg in st.session_state["historico_chat"]:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])

        prompt_ajuste = st.chat_input("Digite aqui o que deseja alterar ou pergunte mais detalhes sobre os projetos...")
        
        if prompt_ajuste:
            st.session_state["historico_chat"].append({"role": "user", "content": prompt_ajuste})
            
            with st.spinner("Atualizando resposta com base no seu feedback..."):
                try:
                    client = genai.Client(api_key=perfil.get("api_key"))
                    
                    contexto_chat = f"BANCO DE DADOS DO USUÁRIO:\n{carregar_base_conhecimento(usuario)}\n\n"
                    for msg in st.session_state["historico_chat"]:
                        contexto_chat += f"\n[{msg['role'].upper()}]: {msg['content']}\n"
                    
                    response_ajuste_text = chamar_gemini_com_fallback(
                        client=client,
                        contents=contexto_chat
                    )

                    st.session_state["historico_chat"].append({"role": "assistant", "content": response_ajuste_text})
                    st.rerun()

                except Exception as e:
                    st.error(f"Erro ao aplicar ajustes: {e}")